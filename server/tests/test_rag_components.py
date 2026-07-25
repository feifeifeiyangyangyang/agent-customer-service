from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

from app.embeddings.mock_embedding import MockEmbeddingClient
from app.rag.chunker import chunk_text
from app.rag.document_parser import extract_text
from app.schemas.retrieval import RetrievalCandidate, RetrievalQueryContext
from app.services.knowledge_service import (
    KnowledgeService,
    _condition_matches,
    _extract_after_sale_type,
    dedupe_candidates,
    heuristic_rerank,
    rrf_fuse,
)


def test_text_parser_decodes_markdown_content() -> None:
    text = extract_text("policy.md", "# 退换货政策\n商品未影响二次销售时可以提交退货申请。".encode())

    assert "退换货政策" in text
    assert "二次销售" in text


def test_docx_parser_extracts_paragraphs(tmp_path: Path) -> None:
    path = tmp_path / "policy.docx"
    document = Document()
    document.add_paragraph("退款会按原支付路径退回。")
    document.add_paragraph("质量问题可以申请换货。")
    document.save(str(path))

    text = extract_text(path.name, path.read_bytes())

    assert "原支付路径" in text
    assert "质量问题" in text


def test_chunk_text_splits_long_policy_with_overlap() -> None:
    text = "发货规则：" + "下单后 24 小时内发货。" * 80

    chunks = chunk_text(text, max_chars=180, overlap=30)

    assert len(chunks) > 1
    assert chunks[0].index == 0
    assert all(chunk.char_count <= 180 for chunk in chunks)
    assert all(chunk.content_hash for chunk in chunks)


def test_mock_embedding_is_deterministic_and_normalized() -> None:
    client = MockEmbeddingClient(32)

    first = client.embed("商品发货时间")
    second = client.embed("商品发货时间")

    assert first == second
    assert len(first) == 32
    assert abs(sum(value * value for value in first) - 1.0) < 0.000001


def test_rrf_fusion_deduplicates_same_chunk() -> None:
    keyword = RetrievalCandidate(
        candidate_id="chunk:1",
        source_type="keyword",
        content="拆封后未影响二次销售可以退货",
        document_id="1",
        chunk_id="1",
        rule_id=None,
        metadata={"matched_terms": ["拆封", "退货"]},
        original_score=0.7,
    )
    dense = keyword.model_copy(update={"source_type": "dense_vector", "original_score": 0.9})
    other = RetrievalCandidate(
        candidate_id="chunk:2",
        source_type="keyword",
        content="发货规则",
        document_id="1",
        chunk_id="2",
        rule_id=None,
        metadata={},
        original_score=0.4,
    )

    fused = rrf_fuse([[keyword, other], [dense]])

    assert [item.candidate_id for item in fused].count("chunk:1") == 1
    assert fused[0].candidate_id == "chunk:1"
    assert fused[0].fused_score is not None


def test_heuristic_rerank_promotes_structured_rule_for_order_context() -> None:
    doc = RetrievalCandidate(
        candidate_id="chunk:1",
        source_type="keyword",
        content="通用文档：拆封可能需要人工判断",
        document_id="1",
        chunk_id="1",
        rule_id=None,
        metadata={},
        original_score=0.8,
        fused_score=0.01,
    )
    rule = RetrievalCandidate(
        candidate_id="rule:1",
        source_type="structured_rule",
        content="结构化规则：签收七天内，未影响二次销售可退货",
        document_id=None,
        chunk_id=None,
        rule_id="1",
        metadata={"rule_version": "AS-2026-07"},
        original_score=0.6,
        fused_score=0.01,
    )

    reranked = heuristic_rerank(
        "这个拆封后还能退吗",
        [doc, rule],
        RetrievalQueryContext(has_specific_order=True, order_status="SIGNED", after_sale_type="RETURN"),
    )

    assert reranked[0].source_type == "structured_rule"
    assert reranked[0].decision_reason == "命中当前订单状态适用的结构化业务规则"


def test_structured_condition_filters_inapplicable_order_status() -> None:
    condition = SimpleNamespace(
        product_category=None,
        order_status="SIGNED",
        payment_status=None,
        shipment_status=None,
        signed_within_days=7,
        after_sale_type="RETURN",
    )

    matched, _ = _condition_matches(
        condition,  # type: ignore[arg-type]
        RetrievalQueryContext(has_specific_order=True, order_status="WAITING_SHIPMENT", after_sale_type="RETURN"),
    )

    assert matched is False


def test_structured_condition_allows_general_rule_without_specific_order() -> None:
    condition = SimpleNamespace(
        product_category=None,
        order_status="SIGNED",
        payment_status=None,
        shipment_status=None,
        signed_within_days=7,
        after_sale_type="RETURN",
    )

    matched, reasons = _condition_matches(
        condition,  # type: ignore[arg-type]
        RetrievalQueryContext(has_specific_order=False, after_sale_type="RETURN"),
    )

    assert matched is True
    assert "order_status=SIGNED" in reasons


def test_freight_question_routes_to_freight_rule_type() -> None:
    assert _extract_after_sale_type("退货包运费吗") == "RETURN_FREIGHT"
    assert _extract_after_sale_type("退货邮费谁承担") == "RETURN_FREIGHT"


def test_dedupe_candidates_by_candidate_id() -> None:
    first = RetrievalCandidate(
        candidate_id="rule:1",
        source_type="structured_rule",
        content="新规则",
        document_id=None,
        chunk_id=None,
        rule_id="1",
        metadata={},
        original_score=1,
    )
    duplicate = first.model_copy(update={"content": "旧规则"})

    assert dedupe_candidates([first, duplicate]) == [first]


@pytest.mark.asyncio
async def test_dense_vector_recall_returns_empty_when_qdrant_unavailable(monkeypatch: pytest.MonkeyPatch) -> None:
    async def broken_search(_query_vector: list[float], _limit: int) -> list[object]:
        raise RuntimeError("qdrant down")

    from app.repositories.qdrant_store import qdrant_store

    monkeypatch.setattr(qdrant_store, "search", broken_search)

    assert await KnowledgeService().dense_vector_recall("拆封还能退吗", 3) == []


@pytest.mark.asyncio
async def test_retrieve_continues_when_keyword_recall_fails(monkeypatch: pytest.MonkeyPatch) -> None:
    service = KnowledgeService()
    vector_candidate = RetrievalCandidate(
        candidate_id="chunk:9",
        source_type="dense_vector",
        content="语义命中：商品外包装损坏",
        document_id="1",
        chunk_id="9",
        rule_id=None,
        metadata={},
        original_score=0.7,
    )
    rule_candidate = RetrievalCandidate(
        candidate_id="rule:9",
        source_type="structured_rule",
        content="结构化规则：破损商品保留凭证后申请售后",
        document_id=None,
        chunk_id=None,
        rule_id="9",
        metadata={"rule_version": "AS-2026-07"},
        original_score=0.8,
    )

    async def broken_keyword(*_args: object, **_kwargs: object) -> list[RetrievalCandidate]:
        raise RuntimeError("keyword down")

    async def fake_dense(*_args: object, **_kwargs: object) -> list[RetrievalCandidate]:
        return [vector_candidate]

    async def fake_rule(*_args: object, **_kwargs: object) -> list[RetrievalCandidate]:
        return [rule_candidate]

    async def empty_cache(*_args: object, **_kwargs: object) -> None:
        return None

    async def noop_cache(*_args: object, **_kwargs: object) -> None:
        return None

    from app.services.redis_runtime_service import redis_runtime_service

    monkeypatch.setattr(service, "keyword_recall", broken_keyword)
    monkeypatch.setattr(service, "dense_vector_recall", fake_dense)
    monkeypatch.setattr(service, "structured_rule_recall", fake_rule)
    monkeypatch.setattr(redis_runtime_service, "get_json", empty_cache)
    monkeypatch.setattr(redis_runtime_service, "set_json", noop_cache)

    candidates = await service.retrieve(object(), "收到商品坏了怎么办", limit=3)  # type: ignore[arg-type]

    assert {candidate.source_type for candidate in candidates} == {"dense_vector", "structured_rule"}
    assert any(
        diagnostic.channel == "keyword" and diagnostic.status == "FAILED"
        for diagnostic in service.last_diagnostics
    )


@pytest.mark.asyncio
async def test_retrieve_returns_empty_when_all_channels_below_threshold(monkeypatch: pytest.MonkeyPatch) -> None:
    service = KnowledgeService()
    weak = RetrievalCandidate(
        candidate_id="chunk:weak",
        source_type="dense_vector",
        content="完全无关内容",
        document_id="1",
        chunk_id="1",
        rule_id=None,
        metadata={},
        original_score=0.01,
    )

    async def empty_keyword(*_args: object, **_kwargs: object) -> list[RetrievalCandidate]:
        return []

    async def weak_dense(*_args: object, **_kwargs: object) -> list[RetrievalCandidate]:
        return [weak]

    async def empty_rule(*_args: object, **_kwargs: object) -> list[RetrievalCandidate]:
        return []

    monkeypatch.setattr(service, "keyword_recall", empty_keyword)
    monkeypatch.setattr(service, "dense_vector_recall", weak_dense)
    monkeypatch.setattr(service, "structured_rule_recall", empty_rule)

    candidates = await service.retrieve(object(), "天上的云是什么味道", limit=3)  # type: ignore[arg-type]

    assert candidates == []
